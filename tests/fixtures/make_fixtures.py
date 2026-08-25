"""
Generates the small binary fixtures used by the test suite
(tests/fixtures/sample.pdf, sample.docx, sample.xlsx, questions.csv).

Run once to (re)build them:
    python tests/fixtures/make_fixtures.py

The files are checked into git (they're tiny and deterministic) so tests
don't need to regenerate them on every run — but this script is the source
of truth if they ever need to change.
"""

from __future__ import annotations

from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from openpyxl import Workbook

FIXTURES_DIR = Path(__file__).parent

PDF_MARKER = "PDF_FIXTURE_MARKER: the quarterly report shows revenue of $42,000."
DOCX_MARKER = "DOCX_FIXTURE_MARKER: the merger agreement was signed in March."
XLSX_MARKER = "XLSX_FIXTURE_MARKER"


def make_pdf() -> None:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), PDF_MARKER)
    doc.save(FIXTURES_DIR / "sample.pdf")
    doc.close()


def make_docx() -> None:
    doc = DocxDocument()
    doc.add_paragraph(DOCX_MARKER)
    doc.add_paragraph("A second paragraph with unrelated filler text.")
    doc.save(FIXTURES_DIR / "sample.docx")


def make_xlsx() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["marker", "value"])
    ws.append([XLSX_MARKER, 42])
    wb.save(FIXTURES_DIR / "sample.xlsx")


def make_questions_csv() -> None:
    (FIXTURES_DIR / "questions.csv").write_text(
        "question\n"
        "What was the quarterly revenue?\n"
        "When was the merger agreement signed?\n"
        "How many documents are in the set?\n"
    )


if __name__ == "__main__":
    make_pdf()
    make_docx()
    make_xlsx()
    make_questions_csv()
    print(f"Fixtures written to {FIXTURES_DIR}")
